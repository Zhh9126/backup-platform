# -*- coding: utf-8 -*-
"""
报告生成器：支持 CSV / Word(docx) / PDF 三种格式。

依赖：python-docx（Word）、reportlab（PDF）
"""
import io
import csv
import time
from datetime import datetime


# ----------------- CSV -----------------
def build_csv(title: str, headers: list, rows: list) -> bytes:
    """生成 UTF-8 BOM 的 CSV 字节流（Excel 直接打开不乱码）。"""
    buf = io.StringIO()
    w = csv.writer(buf)
    w.writerow([title])
    w.writerow(["生成时间", datetime.now().strftime("%Y-%m-%d %H:%M:%S")])
    w.writerow([])
    w.writerow(headers)
    for r in rows:
        w.writerow(r)
    return buf.getvalue().encode("utf-8-sig")


# ----------------- 通用文本清洗 -----------------
_CTRL_RE = None

def _cell_text(val, max_len: int = 0) -> str:
    """把任意值转为报告可安全渲染的文本。

    - 剥离 XML/Word 不接受的控制字符（\\x00-\\x08、\\x0b、\\x0c、\\x0e-\\x1f）；
    - 超长内容按 max_len 截断（报告不是全量日志，过长会撑爆版面）。
    """
    global _CTRL_RE
    if _CTRL_RE is None:
        import re as _re
        _CTRL_RE = _re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f]")
    text = _CTRL_RE.sub("", "" if val is None else str(val))
    if max_len and len(text) > max_len:
        text = text[:max_len] + "…(截断)"
    return text


# ----------------- Word (.docx) -----------------
def build_docx(title: str, summary: dict, headers: list, rows: list) -> bytes:
    """生成 Word 报告（含标题/摘要/表格）。"""
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL

    doc = Document()
    # 标题
    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 元信息
    meta = doc.add_paragraph()
    meta_run = meta.add_run(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    if summary:
        doc.add_paragraph()
        # 摘要小节
        sh = doc.add_heading("汇总信息", level=1)
        for k, v in summary.items():
            p = doc.add_paragraph()
            r1 = p.add_run(f"  {k}：")
            r1.bold = True
            p.add_run(str(v))
    # 表格
    doc.add_paragraph()
    doc.add_heading("详细数据", level=1)
    if rows:
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = "Light Grid Accent 1"
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 表头
        hdr = table.rows[0]
        for i, h in enumerate(headers):
            cell = hdr.cells[i]
            cell.text = str(h)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True
                    r.font.size = Pt(10)
        # 数据
        for ri, row in enumerate(rows, start=1):
            for ci, val in enumerate(row):
                cell = table.rows[ri].cells[ci]
                cell.text = _cell_text(val, max_len=1000)
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(9)
    else:
        doc.add_paragraph("暂无数据")
    # 落盘为字节
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ----------------- PDF -----------------
def build_pdf(title: str, summary: dict, headers: list, rows: list) -> bytes:
    """生成 PDF 报告。reportlab 引擎。"""
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                     Table, TableStyle, PageBreak)
    from reportlab.lib.enums import TA_CENTER

    buf = io.BytesIO()
    # 横版 A4 容纳更多列
    doc = SimpleDocTemplate(
        buf, pagesize=landscape(A4),
        leftMargin=1.5*cm, rightMargin=1.5*cm,
        topMargin=1.5*cm, bottomMargin=1.5*cm,
        title=title, author="Backup Management Platform",
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("title", parent=styles["Title"],
                                  fontSize=18, alignment=TA_CENTER,
                                  textColor=colors.HexColor("#0D9488"))
    sub_style = ParagraphStyle("sub", parent=styles["Normal"],
                                fontSize=9, alignment=TA_CENTER,
                                textColor=colors.grey)
    h1_style = ParagraphStyle("h1", parent=styles["Heading2"],
                               textColor=colors.HexColor("#0F172A"))

    story = []
    story.append(Paragraph(title, title_style))
    story.append(Paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", sub_style))
    story.append(Spacer(1, 0.5*cm))

    if summary:
        story.append(Paragraph("汇总信息", h1_style))
        # 摘要渲染为 2 列
        sum_data = []
        keys = list(summary.items())
        for i in range(0, len(keys), 2):
            row = []
            for j in range(2):
                if i + j < len(keys):
                    k, v = keys[i + j]
                    row.append(f"<b>{k}</b>：{v}")
                else:
                    row.append("")
            sum_data.append(row)
        t = Table(sum_data, colWidths=[12*cm, 12*cm])
        t.setStyle(TableStyle([
            ("FONTSIZE", (0, 0), (-1, -1), 10),
            ("TEXTCOLOR", (0, 0), (-1, -1), colors.HexColor("#0F172A")),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]))
        story.append(t)
        story.append(Spacer(1, 0.5*cm))

    story.append(Paragraph("详细数据", h1_style))
    if rows:
        cell_style = ParagraphStyle("td", parent=styles["Normal"], fontSize=8)
        th_style = ParagraphStyle("th", parent=styles["Normal"], fontSize=9, textColor=colors.white)
        # 数据：首行是表头（teal 底），后续是数据（清洗控制字符 + 截断超长：
        # PDF 表格单行不可跨页，列窄时长文本会把行高撑过页高导致 LayoutError）
        data = [[Paragraph(_cell_text(h, 100), th_style) for h in headers]]
        for row in rows:
            data.append([Paragraph(_cell_text(v, 150), cell_style) for v in row])
        # 自适应列宽：按可用页宽均分（13 列以内可完整容纳，不溢出版面）
        avail = landscape(A4)[0] - 3*cm  # 横向 A4 减左右页边距
        col_w = max(1.6*cm, min(8*cm, avail / max(len(headers), 1)))
        t = Table(data, colWidths=[col_w]*len(headers), repeatRows=1)
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0D9488")),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("GRID", (0, 0), (-1, -1), 0.3, colors.grey),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ]))
        story.append(t)
    else:
        story.append(Paragraph("暂无数据", styles["Normal"]))

    # 页脚
    def _on_page(canvas, doc):
        canvas.saveState()
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(colors.grey)
        canvas.drawString(1.5*cm, 0.8*cm, f"数据备份管理平台 — 报告生成于 {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        canvas.drawRightString(landscape(A4)[0] - 1.5*cm, 0.8*cm, f"第 {doc.page} 页")
        canvas.restoreState()

    doc.build(story, onFirstPage=_on_page, onLaterPages=_on_page)
    return buf.getvalue()


# ----------------- 工厂 -----------------
def build_report(fmt: str, title: str, summary: dict, headers: list, rows: list) -> tuple:
    """根据 fmt 返回 (mime_type, file_ext, content_bytes)。

    可选依赖缺失时抛 ValueError（API 层转 400 明确提示），不产生 500。
    """
    fmt = (fmt or "csv").lower()
    if fmt == "csv":
        return ("text/csv; charset=utf-8-sig", "csv", build_csv(title, headers, rows))
    if fmt in ("docx", "word"):
        try:
            content = build_docx(title, summary, headers, rows)
        except ImportError:
            raise ValueError(
                "Word 导出依赖 python-docx 未安装：请在平台环境执行 "
                "`pip install python-docx` 后重试（CSV 导出不受影响）")
        return ("application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "docx", content)
    if fmt == "pdf":
        try:
            content = build_pdf(title, summary, headers, rows)
        except ImportError:
            raise ValueError(
                "PDF 导出依赖 reportlab 未安装：请在平台环境执行 "
                "`pip install reportlab` 后重试（CSV 导出不受影响）")
        return ("application/pdf", "pdf", content)
    raise ValueError(f"不支持的格式: {fmt}")
